import pandas as pd
import streamlit as st
import numpy as np
from io import BytesIO  # 🟢 Corregido: Importación necesaria para reportes
from scipy.optimize import brentq
from scipy.interpolate import CubicSpline

# Configuración de interfaz en pantalla ancha
st.set_page_config(layout="wide", page_title="Diseño de Pavimento Flexible MCV3")
# 2. Título visible dentro de la aplicación
st.title("Diseño de Pavimentos Flexibles MCV3")
st.markdown("---") # Línea divisoria opcional

# 🟢 IMPORTACIÓN COMPLETA DESDE TU ARCHIVO EXTERNO
#from motores_calculo import (
#    get_mr, calcular_ne_mcv3, calcular_factor_confiabilidad, 
#    calcular_transito_diseno, validar_espesor_mop, generar_excel_vial,
#    calcular_sn_requerido_aashto, calcular_sn_requerido_mop
#)

#------------------------------------------------------------------------------
#------------- MOTORES DE CÁLCULO ---------------------------------------------
#------------------------------------------------------------------------------
        
def get_mr(cbr, tipo_capa="subrasante"):
    if cbr is None or cbr < 2: 
        return "Error, CBR < 2"
    if tipo_capa == "subrasante":
        if cbr < 12: return round(17.6 * (cbr ** 0.64), 2)
        if cbr <= 80: return round(22.1 * (cbr ** 0.55), 2)
        return "Error, CBR > 80"
    elif tipo_capa == "base":
        if cbr == 80.0: return 246.09
        return round(10.5 * (cbr ** 0.72), 2)
    elif tipo_capa == "subbase":
        if cbr == 40.0: return 168.08
        return round(13.2 * (cbr ** 0.69), 2)
    return "Capa no reconocida"

def calcular_factor_confiabilidad(confianza, so):
    from scipy.stats import norm
    zr = norm.ppf(1 - confianza / 100)
    fr = 10 ** (-zr * so)
    return round(fr, 3), zr

def calcular_transito_diseno(trafico_solicitante, fr):
    return round(trafico_solicitante * fr, 0)

def validar_espesor_mop(capa, espesor_mm, trafico_ee):
    if espesor_mm == 0:
        return True, 0, "No considerada"
    if "Concreto asfáltico" in capa: # or "intermedia" in capa or "Base asfáltica" in capa:
#        if trafico_ee < 1_000_000: min_req = 50
#        elif trafico_ee < 5_000_000: min_req = 60
#        else: min_req = 80
        min_req = 50
    elif "Base granular" in capa: min_req = 150
    elif "Subbase granular" in capa: min_req = 120
    else: return True, 0, "OK"
    cumple = espesor_mm >= min_req
    return cumple, min_req, "OK" if cumple else f"⚠️ Mínimo req: {min_req} mm"

# ==========================================================
# 🟢 LAS DOS FUNCIONES DE OPTIMIZACIÓN EN PARALELO
# ==========================================================
def ecuacion_aashto_error(SN, W18, Zr, So, delta_PSI, MR_psi):
    if SN <= 0: 
        return 1e6
    factor_serviciabilidad = np.log10(delta_PSI / (4.2 - 1.5)) / (0.40 + (1094 / (SN + 1) ** 5.19))
    return (Zr * So + 9.36 * np.log10(SN + 1) - 0.20 + factor_serviciabilidad + 2.32 * np.log10(MR_psi) - 8.07 - np.log10(W18))

def calcular_sn_requerido_aashto(W18, Zr, So, delta_PSI, MR_MPa):
    MR_psi = MR_MPa * 145.038
    try:
        sn_pulgadas = brentq(ecuacion_aashto_error, 0.1, 15.0, args=(W18, Zr, So, delta_PSI, MR_psi))
        return round(sn_pulgadas*25.4, 2)
    except ValueError:
        return 0.0

def ecuacion_mop_error(ne, Td, Mr, pi, pf, zr, so):
    if ne <= 0: return -Td
    exponente_serviciabilidad = 1 / (0.4 + (97.81 / (ne + 25.4)) ** 5.19)
    term_serviciabilidad = ((pi - pf) / (pi - 1.5)) ** exponente_serviciabilidad
    Tto = ((ne + 25.4) ** 9.36) * (10 ** (-16.4 + zr * so)) * (Mr ** 2.32) * term_serviciabilidad
    return Tto - Td

def calcular_sn_requerido_mop(Td, zr, so, Mr_MPa, pi, pf):
    try:
        ne_mm = brentq(ecuacion_mop_error, 1.0, 500.0, args=(Td, Mr_MPa, pi, pf, zr, so))
        return round(ne_mm, 2)
    except ValueError:
        return 0.0

def calcular_ne_mcv3(ee, mpa, tmapa):
    if ee <= 0: return 0.0
    
    # Eje X en escala logarítmica
    ee_x_log = np.log10(np.array([1.0, 5.0, 10.0, 20.0, 50.0, 100.0]))
    ee_val_log = np.log10(max(ee, 1.0))
    
    # Datos corregidos (mismas claves 36, 50, 60, 80, 100, 120, 140, 200)
    p6 = {
        36: [37.0, 50.0, 55.0, 60.0, 68.0, 74.5], 50: [34.5, 47.0, 52.0, 57.0, 65.0, 72.0],
        60: [33.0, 45.0, 50.0, 55.0, 63.5, 70.0], 80: [30.5, 42.5, 47.0, 52.0, 60.5, 68.5],
        100: [28.5, 40.0, 45.5, 50.0, 58.0, 67.0], 120: [26.5, 38.0, 43.0, 47.5, 56.0, 64.5],
        140: [24.5, 36.0, 41.5, 46.0, 55.0, 63.5], 200: [22.5, 26.5, 35.0, 41.0, 50.0, 57.0]
    }
    p14 = {
        36: [45.0, 57.0, 63.0, 68.0, 80.0, 91.0], 50: [41.0, 52.0, 59.0, 64.0, 75.0, 88.0],
        60: [39.0, 50.0, 56.0, 62.0, 72.0, 85.0], 80: [36.0, 46.0, 53.0, 58.0, 69.0, 80.0],
        100: [32.0, 43.0, 49.0, 54.0, 65.0, 79.0], 120: [27.0, 40.0, 47.0, 51.0, 62.0, 78.0],
        140: [25.0, 37.0, 43.0, 48.0, 59.0, 75.0], 200: [22.5, 24.0, 34.0, 41.0, 51.0, 64.0]
    }
    p19 = {
        36: [53.0, 67.0, 74.0, 80.0, 93.0, 106.0], 50: [47.5, 61.0, 69.0, 75.0, 88.0, 100.0], # Clave 50 corregida
        60: [43.0, 57.0, 66.0, 71.0, 84.0, 96.0], 80: [40.0, 53.5, 61.5, 67.5, 80.5, 92.0],
        100: [36.0, 49.5, 57.5, 63.5, 74.5, 86.0], 120: [31.5, 44.0, 51.0, 57.0, 68.0, 76.0],
        140: [28.5, 40.5, 47.5, 51.5, 62.5, 70.5], 200: [21.5, 22.5, 27.0, 35.0, 45.0, 55.0] # Array 200 corregido
    }

    def get_vals(d, ee_log, m):
        ks = sorted(d.keys())
        m_c = max(min(m, max(ks)), min(ks))
        vals_en_ee = [float(CubicSpline(ee_x_log, d[k], bc_type='natural')(ee_log)) for k in ks]
        return np.interp(m_c, ks, vals_en_ee)

    val_6 = get_vals(p6, ee_val_log, mpa)
    val_14 = get_vals(p14, ee_val_log, mpa)
    val_19 = get_vals(p19, ee_val_log, mpa)

    resultado = np.interp(max(min(tmapa, 19.0), 6.0), [6.0, 14.0, 19.0], [val_6, val_14, val_19])
    return round(float(resultado), 2)


# Títulos superiores del proyecto
with st.expander("📝 Identificación y Datos Generales del Proyecto", expanded=True):
    c_cam, c_sec = st.columns(2)
    nombre_camino = c_cam.text_input("Nombre del Camino:", value="Ruta 21-CH")
    nombre_sector = c_sec.text_input("Sector / Tramo:", value="Cebollar-Ollagüe")

# --- 1. CARGA DE BASE DE DATOS (ARRAYS NATIVOS) ---
@st.cache_data
def cargar_base_clima():
    estaciones = np.array(['AEROPUERTO CHACALLUTA', 'AZAPA', 'CAQUENA', 'CHILCAYA', 'CHUNGARA AJATA', 'CODPA', 'EL BUITRE AERÓDROMO', 'PARINACOTA EX ENDESA', 'PUTRE', 'AEROPUERTO EN IQUIQUE', 'CERRO COLORADO', 'COYACAGUA', 'HUARA EN FUERTE BAQUEDANO', 'LAGUNILLAS (PAMPA LIRIMA)', 'AEROPUERTO EL LOA (CALAMA)', 'AEROPUERTO CERRO MORENO -ANTOF.', 'AGUAS VERDES', 'AYQUINA', 'CASPANA', 'CHIU-CHIU', 'CONCHI EMBALSE', 'EL TATIO *', 'LEQUENA', 'LINZOR', 'OLLAGUE', 'PARSHALL N2', 'PEINE', 'SIERRA GORDA', 'TOCONAO EXPERIMENTAL', 'TOCONCE', 'AEROPUERTO D. ATACAMA', 'CANTO DE AGUA', 'CONAY EN ALBARICOQUE', 'EL TRANSITO', 'HUASCO BAJO', 'IGLESIA COLORADA', 'LA COM PAÑ ÍA (VAL LENAR )', 'LAUTARO EMBALSE', 'LOS LOROS', 'PORTEZUELO EL GAUCHO', 'SAN FELIX', 'SANTAJUANA', 'AEROPUERTO LA SERENA', 'CAIMANES', 'CAREN', 'CERRO VEGA NEGRA', 'COGOTI EMBALSE', 'EL SOLDADO', 'EL TRAPICHE', 'HURTADO', 'ILLAPEL DGA', 'JUNTAS', 'LA LAGUNA EMBALSE', 'LA ORTIGA', 'LA TRANQUILLA', 'LAS RAMADAS', 'LOS CONDORES', 'PALOMA EMBALSE', 'RECOLETA EMBALSE', 'RIVADAVIA', 'AEROPUERTO RODELILLO', 'AEROPUERTO STO. DOMINGO', 'ALICAHUE', 'LAGO PEÑUELAS', 'LLIU-LLIU EMBALSE', 'LOS AROMOS', 'PORTILLO*', 'QUILLOTA', 'RIO ACONCAGUA EN CHACABUQUITO', 'VILCUYA', 'AEROPUERTO QTA. NORMAL', 'AEROPUERTO TOBALABA', 'AEROPUERTO PUDAHUEL', 'CERRO CALAN', 'EL YESO EMBALSE', 'HUECHUN ANDINA', 'LAGUNA ACULEO', 'LAGUNA NEGRA', 'LOS PANGUILES', 'MELIPILLA', 'PIRQUE', 'QUEBRADA DE MACUL', 'RIO MAPOCHO EN LOS ALMENDROS', 'RUNGUE EMBALSE', 'CONVENTO VIEJO', 'RENGO', 'TERMAS DEL FLACO', 'AEROPUERTO CURICÓ', 'ANCOA EMBALSE', 'COLORADO', 'DIGUA AMBALSE', 'LO AGUIRRE', 'PARRAL', 'PENCAHUE', 'POTRERO GRANDE', 'RIO CLARO EN RAUQUEN', 'RIO MAULE EN ARMERILLO', 'RIO MELADO EN EL SALTO', 'TALCA U.C', 'AEROPUERTO CHILLÁN', 'AEROPUERTO CONCEPCIÓN', 'AEROPUERTO LOS ANGELES', 'ALTO MALLINES', 'CARACOL', 'COIHUECO EMBALSE', 'CONTULMO', 'DIGUILLIN', 'PARQUE NAHUELBUTA', 'QUILACO', 'RIO BIO-BIO EN LLANQUEN', 'AEROPUERTO TEMUCO', 'ANGOL', 'CHERQUENCO', 'ERCILLA (VIDA NUEVA)', 'LAGUNA MALLECO', 'LAUTARO EMBALSE', 'LIUCURA', 'LONQUIMAY', 'MALANCAHUELLO', 'PUCON', 'PUEBLO NUEVO (TEMUCO)', 'PUERTO SAAVEDRA', 'PUESCO (ADUANA)', 'TEODORO SCHMIT', 'TRAIGUEN', 'TRICAUCO', 'AEROPUERTO VALDIVIA', 'LAGO RANCO', 'RIO FUI EN DASAGUE LAGO PIRIHUICO', 'VALDIVIA (U. AUSTRAL)', 'Alto Pa lena Ad,', 'Chaitén, Ad,', 'Futa leufú Ad,', 'Cañal Bajo, Osorno Ad,', 'El Tepual Puerto Montt Ap,', 'ADOLFO MATTHEI', 'Puerto Aysén Ad,', 'Lord Cochrane Ad,', 'BAHIA MURTA', 'CHILE CHICO', 'COYHAIQUE (ESCUELA AGRICOLA)', 'COYHAIQUE ALTO', 'COYHAIQUE CONAF', 'ESTANCIA BAÑO NUEVO', 'LAJUNTA', 'LAGO VERDE', 'ÑIREHUAO', 'PUERTO CISNES', 'PUERTO GUADAL', 'PUERTO IBAÑEZ', 'PUERTO PUYUHUAPI', 'RIO BAKER EN ANGOSTURA CHACABUCO', 'RIO CISNES', 'RIO PASCUA ANTE JUNTA RIO QUETRU', 'VILLA MAÑIHUALES', 'VILLA OHIGGINS', 'VILLA ORTEGA', 'BAHIA SAN FELIPE', 'CERRO CASTILLO', 'CERRO GUIDO', 'ISLA RIESCO', 'MONTE AYMOND', 'ONAISIN EN MARIA CRISTINA', 'PAMPAHUANACO', 'PORVENIR', 'PUERTO NATALES', 'PUNTA ARENAS', 'RIO LAS CHINAS EN CERRO GUIDO', 'RUSSFIN', 'SAN SEBASTIAN', 'VILLA TEHUELCHE', 'TORRES DEL PAINE'])
    tmpa = np.array([19.6, 19.1, 3.5, 3.4, 4.2, 16.1, 19.8, 2.6, 8.9, 18.7, 15.5, 5.7, 17.6, 3.2, 13.4, 17.1, 14.4, 12.4, 17.1, 11.9, 9.5, 2.3, 9.0, 5.0, 9.3, 9.4, 17.3, 18.5, 17.9, 11.1, 16.2, 17.1, 17.7, 18.0, 15.2, 18.8, 15.9, 19.3, 18.4, 2.0, 19.1, 18.4, 14.7, 14.8, 17.6, 3.6, 18.3, 3.1, 16.7, 17.8, 16.7, 14.9, 9.6, 16.9, 17.7, 17.3, 16.3, 17.7, 17.0, 18.1, 14.7, 13.1, 15.6, 13.8, 16.1, 15.7, 6.2, 15.2, 16.7, 16.5, 17.1, 17.0, 16.6, 17.8, 10.5, 17.3, 15.9, 7.2, 16.5, 15.5, 15.0, 15.6, 15.6, 16.8, 15.9, 16.3, 8.1, 16.0, 15.0, 14.6, 13.9, 9.4, 15.7, 16.5, 14.3, 17.3, 16.1, 14.9, 16.0, 15.0, 13.5, 15.0, 9.4, 14.0, 14.9, 14.1, 13.4, 10.1, 14.6, 11.8, 12.9, 15.5, 10.7, 13.8, 10.1, 13.8, 10.4, 10.3, 9.5, 13.3, 12.8, 12.5, 11.2, 12.5, 13.4, 11.8, 12.6, 12.4, 10.6, 13.4, 11.0, 11.0, 11.0, 11.2, 12.0, 12.5, 10.1, 9.9, 9.8, 11.1, 8.8, 7.5, 8.9, 7.7, 10.6, 9.6, 8.5, 10.0, 10.1, 10.5, 10.0, 10.1, 8.0, 9.1, 10.0, 8.8, 9.6, 7.2, 8.3, 9.0, 7.6, 7.3, 7.0, 5.4, 7.4, 7.5, 7.6, 9.3, 4.7, 6.7, 6.7, 8.8])
    lat_sur = np.array([18.35, 18.52, 18.05, 18.79, 18.24, 18.83, 18.51, 18.2, 18.199, 20.55, 20.07, 20.05, 20.13, 19.9, 22.5, 23.45, 25.4, 22.28, 22.34, 22.34, 22.03, 22.37, 21.66, 22.23, 21.22, 21.94, 23.68, 22.89, 23.19, 22.26, 27.26, 28.1, 28.95, 0.0, 28.47, 28.16, 28.58, 27.98, 27.83, 28.62, 28.93, 28.67, 29.92, 31.93, 30.85, 30.9, 30.01, 32.01, 29.37, 30.29, 31.65, 29.98, 30.2, 30.19, 31.9, 31.02, 32.11, 30.7, 30.51, 29.98, 33.04, 33.66, 32.34, 33.15, 33.1, 32.96, 32.84, 32.9, 32.85, 32.86, 33.45, 33.45, 33.54, 33.4, 33.68, 33.08, 33.89, 33.67, 33.44, 33.68, 33.67, 33.5, 33.37, 33.02, 34.77, 34.42, 34.89, 34.97, 35.91, 35.64, 36.26, 35.97, 36.19, 35.37, 35.18, 35.45, 35.71, 35.88, 35.44, 36.59, 36.78, 37.4, 37.16, 36.65, 36.64, 38.02, 36.87, 37.82, 37.69, 38.2, 38.77, 37.78, 38.68, 38.04, 38.22, 27.98, 38.65, 38.45, 38.47, 39.28, 38.71, 38.79, 39.52, 39.03, 38.26, 38.85, 39.65, 40.32, 39.87, 39.81, 43.612, 42.931, 43.189, 40.605, 41.435, 40.588, 45.396, 47.244, 46.462, 46.544, 45.574, 45.48, 45.551, 45.267, 43.971, 44.239, 45.271, 44.728, 46.843, 46.289, 44.323, 47.141, 44.498, 48.159, 45.173, 48.469, 45.372, 52.87, 51.26, 50.899, 52.882, 52.164, 53.306, 54.05, 53.291, 51.734, 53.123, 51.051, 53.759, 53.322, 52.441, 51.184])
    long_oeste = np.array([70.34, 70.18, 69.2, 69.08, 69.18, 69.74, 70.28, 69.27, 69.56, 70.18, 69.27, 68.81, 69.75, 68.8, 69.9, 70.44, 69.96, 68.32, 68.21, 68.64, 68.62, 68.01, 68.66, 68.02, 68.25, 68.52, 68.06, 69.32, 68.0, 68.17, 70.77, 70.78, 70.11, 0.0, 71.19, 69.88, 70.81, 70.0, 70.11, 70.05, 70.46, 70.66, 71.2, 71.14, 70.77, 70.52, 71.09, 70.32, 71.12, 70.7, 71.19, 70.09, 70.04, 70.48, 70.67, 70.59, 71.31, 71.04, 71.1, 70.56, 71.54, 71.61, 70.75, 71.56, 71.21, 71.35, 70.11, 71.21, 70.51, 70.47, 70.68, 70.55, 70.79, 70.54, 70.09, 70.77, 70.88, 70.11, 71.03, 71.2, 70.59, 70.51, 70.45, 70.91, 71.13, 70.87, 70.33, 71.22, 71.3, 71.26, 71.55, 70.57, 71.83, 71.83, 71.1, 71.73, 71.11, 71.02, 71.62, 72.4, 73.06, 72.42, 71.24, 71.4, 71.8, 73.23, 71.64, 72.96, 72.01, 71.3, 72.64, 72.64, 72.0, 72.46, 71.81, 70.0, 71.09, 71.37, 71.58, 71.95, 72.556, 73.39, 73.55, 73.08, 72.65, 71.55, 73.08, 72.47, 71.89, 73.25, 71.8053, 72.8289, 71.8492, 73.0608, 73.0975, 73.107, 72.6639, 72.5861, 72.669, 71.709, 72.029, 71.604, 72.059, 71.529, 72.406, 71.846, 71.709, 72.681, 72.701, 71.934, 72.56, 72.726, 71.306, 73.089, 72.148, 72.559, 71.982, 69.932, 72.327, 72.331, 71.571, 69.608, 69.268, 68.8, 70.37, 72.478, 70.877, 72.517, 69.189, 68.661, 71.402, 72.967])
    alt_msnm = np.array([63, 365, 4400, 4270, 4585, 1870, 110, 4420, 3545, 52, 2510, 4013, 1100, 4020, 2293, 113, 1560, 3031, 3260, 2524, 3010, 4370, 3320, 4100, 3700, 3318, 2460, 1616, 2500, 3310, 204, 330, 1600, 0, 50, 1550, 430, 1110, 940, 4000, 1150, 560, 142, 450, 740, 3600, 740, 3290, 300, 1100, 290, 2150, 3160, 1560, 1000, 1380, 190, 320, 350, 820, 330, 75, 750, 360, 260, 100, 3000, 130, 950, 1100, 527, 650, 480, 848, 2475, 590, 360, 2780, 190, 168, 659, 950, 966, 700, 239, 310, 2650, 225, 421, 420, 390, 2000, 175, 55, 445, 64, 470, 730, 130, 151, 12, 120, 1770, 610, 314, 41, 670, 1177, 231, 750, 92, 113, 500, 262, 894, 1110, 1043, 931, 950, 230, 119, 5, 620, 13, 234, 520, 18, 100, 600, 10, 281, 70, 350, 61, 85, 55, 11, 196, 240, 215, 343, 730, 340, 700, 45, 350, 535, 10, 210, 215, 10, 160, 740, 20, 150, 270, 550, 10, 130, 230, 10, 160, 30, 150, 35, 45, 5, 75, 225, 20, 190, 25])
    region = np.array(['ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'ARICA Y PARINACOTA', 'TARAPACA', 'TARAPACA', 'TARAPACA', 'TARAPACA', 'TARAPACA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ANTOFAGASTA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'ATACAMA', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'COQUIMBO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'VALPARAISO', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', 'METROPOLITANA', "BERNARDO O'HIGGINS", "BERNARDO O'HIGGINS", "BERNARDO O'HIGGINS", 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'MAULE', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'BIOBIO', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'ARAUCANIA', 'DE LOS RIOS', 'DE LOS RIOS', 'DE LOS RIOS', 'DE LOS RIOS', 'DE LOS LAGOS', 'DE LOS LAGOS', 'DE LOS LAGOS', 'DE LOS LAGOS', 'DE LOS LAGOS', 'DE LOS LAGOS', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'AYSEN', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES', 'MAGALLANES'])

    df = pd.DataFrame({
        'ESTACIÓN': estaciones,
        'TMPA': tmpa,
        'LAT_SUR': lat_sur,
        'LONG_OESTE': long_oeste,
        'ALT_MSNM': alt_msnm,
        'REGIÓN': region
    })

    df['ESTACIÓN'] = df['ESTACIÓN'].str.strip()
    df['REGIÓN'] = df['REGIÓN'].str.strip()
    return df

df_clima = cargar_base_clima()

# --- 2. INTERFAZ DE BÚSQUEDA INTERACTIVA ---
st.markdown("### Ubicación y Parámetros Climáticos del Proyecto")

lista_estaciones = df_clima['ESTACIÓN'].tolist()
index_default = lista_estaciones.index("OLLAGUE") if "OLLAGUE" in lista_estaciones else 0

col_est, col_info = st.columns([2, 2])

with col_est:
    estacion_seleccionada = st.selectbox(
        "Seleccione la Estación de Referencia:",
        options=lista_estaciones,
        index=index_default
    )

# Extraer parámetros de la estación elegida
datos_estacion = df_clima[df_clima['ESTACIÓN'] == estacion_seleccionada].iloc[0]

tmapa_proyecto = float(datos_estacion['TMPA'])
region_proyecto = str(datos_estacion['REGIÓN'])
lat_proyecto = float(datos_estacion['LAT_SUR'])
long_proyecto = float(datos_estacion['LONG_OESTE'])
alt_proyecto = float(datos_estacion['ALT_MSNM'])

with col_info:
    st.info(
        f"**Región:** {region_proyecto}\n\n"
        f"**TMPA:** {tmapa_proyecto} °C | **Altitud:** {alt_proyecto} msnm\n\n"
        f"**Coordenadas:** {lat_proyecto}° S, {long_proyecto}° O"
    )

with st.expander("🔍 Ver Base de Datos de Estaciones Climáticas Completa"):
    st.dataframe(
        df_clima,
        hide_index=True,
        width="stretch",
        column_config={
            "ESTACIÓN": st.column_config.TextColumn("Estación", width=220),
            "TMPA": st.column_config.NumberColumn("TMAPA (°C)", format="%.1f", width=100),
            "LAT SUR": st.column_config.NumberColumn("Latitud Sur", format="%.2f", width=100),
            "LONG OESTE": st.column_config.NumberColumn("Longitud Oeste", format="%.2f", width=110),
            "ALT msnm": st.column_config.NumberColumn("Altitud (msnm)", format="%d", width=100),
            "REGIÓN": st.column_config.TextColumn("Región", width=200),
        }
    )

# ==========================================
# 1. TABLA: DATOS DE ENTRADA
# ==========================================
st.markdown("### Datos de Entrada")

if "df_entrada" not in st.session_state:
    datos_entrada = {
        "Datos de entrada": [
            "CBR diseño (%)", "Tráfico solicitante (EE)", "TMAPA (°C)",
            "Nivel de confianza (%)", "Error std combinado (So)",
            "Índice serviciabilidad inicial (pi)", "Índice serviciabilidad final (pf)",
            "Factor de confiabilidad Fr", "Módulo resiliente (MPa)", "Tránsito de diseño (T)"
        ],
        # 🟢 CORRECCIÓN: Inicializamos la columna Subrasante estrictamente como TEXTO (strings)
        "Subrasante": ["20.0", "1,430,000", f"{tmapa_proyecto:.1f}", "60.0", "0.450", "4.20", "2.00", "0.000", "0.0", "0"],
        "Base": [80.0, None, None, None, None, None, None, None, 0.0, None],
        "Subbase": [40.0, None, None, None, None, None, None, None, 0.0, None]
    }

    # Creamos el DataFrame base
    df_inicial = pd.DataFrame(datos_entrada)
    
    # Guardamos el DataFrame en el session_state
    st.session_state.df_entrada = df_inicial

# Encabezado (~38px) + ~35px por cada fila + padding extra, si la altura cambia dinámicamente
altura_entrada = 38 + (len(st.session_state.df_entrada) * 35) + 5

# 🟢 FORZADO ABSOLUTO: Obligamos a que la columna sea tipo String. 
# Esto repara el error incluso si tienes una sesión antigua guardada en caché.
st.session_state.df_entrada["Subrasante"] = st.session_state.df_entrada["Subrasante"].astype(str)
    
st.session_state.df_entrada.loc[2, "Subrasante"] = f"{tmapa_proyecto:.1f}"

tabla_entrada = st.data_editor(
    st.session_state.df_entrada,
    disabled=["Datos de entrada"], #"Subrasante", "Base", "Subbase"], # 🟢 Nota: Al volverlo texto formateado, debe quedar deshabilitada la edición directa en celdas individuales para evitar que el usuario rompa el formato de texto escribiendo letras.
    hide_index=True,
    width="stretch",  #"content",
    height=altura_entrada,
    column_config={
        "Datos de entrada": st.column_config.TextColumn(width=240),
        # 🟢 CORREGIDO: Se añade alignment="center" para centrar los valores en pantalla
        "Subrasante": st.column_config.TextColumn(width=120, alignment="center"),  # 🟢 Cambiado a TextColumn
        # Opcional: También puedes centrar los números de las columnas Base y Subbase si lo deseas
        "Base": st.column_config.NumberColumn(format="%.2f", width=100, alignment="center"),
        "Subbase": st.column_config.NumberColumn(format="%.2f", width=100, alignment="center"),
    },
    key="editor_entrada"
)

# 🔄 CÁLCULO REACTIVO DE VARIABLES VIALES
def limpiar_numero(val):
    """Función auxiliar para convertir el texto formateado de la pantalla a número limpio"""
    if val is None:
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    # Removemos las comas de miles si existen y convertimos a decimal
    texto_limpio = str(val).replace(",", "")
    try:
        return float(texto_limpio)
    except ValueError:
        return 0.0

# [PASO 1]: Extraemos y limpiamos los valores para que Python pueda operar matemáticamente 🟢
cbr_sub = limpiar_numero(tabla_entrada.loc[0, "Subrasante"])
cbr_base = limpiar_numero(tabla_entrada.loc[0, "Base"])
cbr_subbase = limpiar_numero(tabla_entrada.loc[0, "Subbase"])
w18_solicitante = limpiar_numero(tabla_entrada.loc[1, "Subrasante"])
confian_val = limpiar_numero(tabla_entrada.loc[3, "Subrasante"])
so_val = limpiar_numero(tabla_entrada.loc[4, "Subrasante"])

pi = limpiar_numero(tabla_entrada.loc[5, "Subrasante"])              
pf = limpiar_numero(tabla_entrada.loc[6, "Subrasante"])              
delta_PSI = pi - pf  # 🟢 CORREGIDO: Ahora sí operará números flotantes sin error

fr_calculado, zr_val = calcular_factor_confiabilidad(confian_val, so_val)
t_diseno_calculado = calcular_transito_diseno(w18_solicitante, fr_calculado)

mr_sub_calculado = get_mr(cbr_sub, tipo_capa="subrasante")
mr_base_calculado = get_mr(cbr_base, tipo_capa="base")
mr_subbase_calculado = get_mr(cbr_subbase, tipo_capa="subbase")

#if isinstance(mr_sub_calculado, float) and mr_sub_calculado > 0:
#    sn_req_subrasante = calcular_sn_requerido(
#        W18=t_diseno_calculado, Zr=zr_val, So=so_val, delta_PSI=delta_PSI, MR_MPa=mr_sub_calculado
#    )
#else:
#    sn_req_subrasante = 0.0

# [PASO 5]: Verificamos si los datos difieren y forzamos el bloqueo de celdas calculadas 🟢

# Preparamos los textos con el formato exacto que DEBEN tener en pantalla
fmt_cbr_sub = f"{cbr_sub:,.1f}"
fmt_w18 = f"{w18_solicitante:,.0f}"
fmt_tmapa = f"{tmapa_proyecto:,.1f}"
fmt_conf = f"{confian_val:,.1f}"
fmt_so = f"{so_val:,.3f}"
fmt_pi = f"{pi:,.2f}"
fmt_pf = f"{pf:,.2f}"
fmt_fr = f"{fr_calculado:,.3f}"
fmt_mr_sub = f"{mr_sub_calculado:,.1f}"
fmt_t_dis = f"{t_diseno_calculado:,.0f}"

# Verificamos si alguna celda tiene un valor o formato distinto al que le corresponde
necesita_actualizar = (
    # 1. Chequeo de inputs (por si el usuario ingresó un número sin comas o modificó el TMAPA)
    tabla_entrada.loc[0, "Subrasante"] != fmt_cbr_sub or
    tabla_entrada.loc[1, "Subrasante"] != fmt_w18 or
    tabla_entrada.loc[2, "Subrasante"] != fmt_tmapa or # TMAPA forzado por el selectbox
    tabla_entrada.loc[3, "Subrasante"] != fmt_conf or
    tabla_entrada.loc[4, "Subrasante"] != fmt_so or
    tabla_entrada.loc[5, "Subrasante"] != fmt_pi or
    tabla_entrada.loc[6, "Subrasante"] != fmt_pf or
    # 2. Chequeo de celdas calculadas (BLOQUEO VIRTUAL: Si el usuario las altera, saltará True)
    tabla_entrada.loc[7, "Subrasante"] != fmt_fr or
    tabla_entrada.loc[8, "Subrasante"] != fmt_mr_sub or
    tabla_entrada.loc[9, "Subrasante"] != fmt_t_dis or
    tabla_entrada.loc[8, "Base"] != mr_base_calculado or
    tabla_entrada.loc[8, "Subbase"] != mr_subbase_calculado
)

# 3. Protegemos que no se escriba accidentalmente en las celdas que deben ir vacías
for fila in [1, 2, 3, 4, 5, 6, 7, 9]:
    if pd.notna(tabla_entrada.loc[fila, "Base"]) or pd.notna(tabla_entrada.loc[fila, "Subbase"]):
        necesita_actualizar = True
        break

# Si algo se desvió, inyectamos la matriz perfecta y reiniciamos
if necesita_actualizar:
    # Restauramos inputs con su formato
    st.session_state.df_entrada.loc[0, "Subrasante"] = fmt_cbr_sub
    st.session_state.df_entrada.loc[1, "Subrasante"] = fmt_w18
    st.session_state.df_entrada.loc[2, "Subrasante"] = fmt_tmapa
    st.session_state.df_entrada.loc[3, "Subrasante"] = fmt_conf
    st.session_state.df_entrada.loc[4, "Subrasante"] = fmt_so
    st.session_state.df_entrada.loc[5, "Subrasante"] = fmt_pi
    st.session_state.df_entrada.loc[6, "Subrasante"] = fmt_pf
    
    # Restauramos/Bloqueamos celdas de cálculo
    st.session_state.df_entrada.loc[7, "Subrasante"] = fmt_fr
    st.session_state.df_entrada.loc[8, "Subrasante"] = fmt_mr_sub
    st.session_state.df_entrada.loc[9, "Subrasante"] = fmt_t_dis
    st.session_state.df_entrada.loc[8, "Base"] = mr_base_calculado
    st.session_state.df_entrada.loc[8, "Subbase"] = mr_subbase_calculado
    
    # Limpiamos celdas vacías por si el usuario tecleó ahí
    for fila in [1, 2, 3, 4, 5, 6, 7, 9]:
        st.session_state.df_entrada.loc[fila, "Base"] = None
        st.session_state.df_entrada.loc[fila, "Subbase"] = None
        
    st.rerun()  # Se ejecutará una sola vez y se detendrá de forma segura

# ==========================================
# 2. TABLA: MATERIALES
# ==========================================
st.markdown("### Características de Materiales")

cbr_base_dinamico = tabla_entrada.loc[0, "Base"] if tabla_entrada.loc[0, "Base"] is not None else 80.0
cbr_subbase_dinamico = tabla_entrada.loc[0, "Subbase"] if tabla_entrada.loc[0, "Subbase"] is not None else 40.0

if "df_materiales" not in st.session_state:
    filas_materiales = [
        {"Capa": "Microaglomerado", "Estabilidad Marshall (N)": 7800.0, "CBR (%)": None, "Coef. Estr. (a)": 0.0},
        {"Capa": "Concreto asfáltico", "Estabilidad Marshall (N)": 9000.0, "CBR (%)": None, "Coef. Estr. (a)": 0.0},
        {"Capa": "Capa intermedia (binder)", "Estabilidad Marshall (N)": 8000.0, "CBR (%)": None, "Coef. Estr. (a)": 0.0},
        {"Capa": "Base asfáltica", "Estabilidad Marshall (N)": 6000.0, "CBR (%)": None, "Coef. Estr. (a)": 0.0},
        {"Capa": "Base granular chancada", "Estabilidad Marshall (N)": None, "CBR (%)": cbr_base_dinamico, "Coef. Estr. (a)": 0.0},
        {"Capa": "Subbase granular", "Estabilidad Marshall (N)": None, "CBR (%)": cbr_subbase_dinamico, "Coef. Estr. (a)": 0.0},
    ]
    st.session_state.df_materiales = pd.DataFrame(filas_materiales)

st.session_state.df_materiales.loc[4, "CBR (%)"] = cbr_base_dinamico
st.session_state.df_materiales.loc[5, "CBR (%)"] = cbr_subbase_dinamico

# Encabezado (~38px) + ~35px por cada fila + padding extra, si la altura cambia dinámicamente
altura_materiales = 38 + (len(st.session_state.df_materiales) * 35) + 5

tabla_materiales = st.data_editor(
    st.session_state.df_materiales,
    disabled=["Capa", "Coef. Estr. (a)"],
    hide_index=True,
    width="content",
    height=altura_materiales,
    column_config={
        "Capa": st.column_config.TextColumn(width=220),
        "Estabilidad Marshall (N)": st.column_config.NumberColumn(width=160, format="%d", min_value=0),
        "CBR (%)": st.column_config.NumberColumn(width=100, format="%.1f", min_value=0),
        "Coef. Estr. (a)": st.column_config.NumberColumn(width=120, format="%.3f"),
    },
    key="editor_materiales"
)

coef_calculados = []
for idx, fila in tabla_materiales.iterrows():
    capa = fila["Capa"]
    marshall = fila["Estabilidad Marshall (N)"]
    cbr = fila["CBR (%)"]
    
    if marshall is not None and marshall > 0:
        a_i = round(0.0078 * (marshall ** 0.441), 3)
    elif capa == "Base granular chancada" and cbr is not None:
        a_i = round(0.032 * (cbr ** 0.32), 3)
    elif capa == "Subbase granular" and cbr is not None:
        a_i = round(0.058 * (cbr ** 0.19), 3)
    else:
        a_i = 0.0
    coef_calculados.append(a_i)

if not tabla_materiales["Coef. Estr. (a)"].equals(pd.Series(coef_calculados)):
    st.session_state.df_materiales["Estabilidad Marshall (N)"] = tabla_materiales["Estabilidad Marshall (N)"]
    st.session_state.df_materiales["CBR (%)"] = tabla_materiales["CBR (%)"]
    st.session_state.df_materiales["Coef. Estr. (a)"] = coef_calculados
    st.rerun()

# ==========================================
# 3. TABLA: ESTRUCTURACIÓN (CORRECCIÓN TRÁFICO)
# ==========================================
st.markdown("### Estructuración y Espesores Mínimos")

# 🟢 MANTENER: Inicialización segura de datos de partida
if "df_estructura" not in st.session_state:
    filas_estructuracion = [
        {"Material / Capa": "Microaglomerado", "Coef. Estr.": 0.40, "Espesor (mm)": 0, "Coef. Drenaje": None, "Total (mm)": 0.0},
        {"Material / Capa": "Concreto asfáltico", "Coef. Estr.": 0.43, "Espesor (mm)": 60, "Coef. Drenaje": None, "Total (mm)": 0.0},
        {"Material / Capa": "Capa intermedia (binder)", "Coef. Estr.": 0.41, "Espesor (mm)": 0, "Coef. Drenaje": None, "Total (mm)": 0.0},
        {"Material / Capa": "Base asfáltica", "Coef. Estr.": 0.36, "Espesor (mm)": 0, "Coef. Drenaje": None, "Total (mm)": 0.0},
        {"Material / Capa": "Base granular chancada", "Coef. Estr.": 0.13, "Espesor (mm)": 150, "Coef. Drenaje": 1.4, "Total (mm)": 0.0},
        {"Material / Capa": "Subbase granular", "Coef. Estr.": 0.12, "Espesor (mm)": 0, "Coef. Drenaje": 1.4, "Total (mm)": 0.0},
    ]
    st.session_state.df_estructura = pd.DataFrame(filas_estructuracion)

# Encabezado (~38px) + ~35px por cada fila + padding extra, si la altura cambia dinámicamente
altura_estructura = 38 + (len(st.session_state.df_estructura) * 35) + 5

# MANTENER: Sincronizamos coeficientes de materiales directo a la estructura
st.session_state.df_estructura["Coef. Estr."] = st.session_state.df_materiales["Coef. Estr. (a)"]

# 🟢 CORRECCIÓN AQUÍ: Purificamos el valor de la pantalla convirtiéndolo a número limpio
trafico_ee_actual = limpiar_numero(tabla_entrada.loc[1, "Subrasante"])

# 🟢 CORRECCIÓN AQUÍ: Un solo bucle for limpio y bien indentado
alertas_espesores = []
for idx, fila in st.session_state.df_estructura.iterrows():
    capa = fila["Material / Capa"]
    espesor = fila["Espesor (mm)"]
    # Ahora validar_espesor_mop recibe un número y no un texto
    _, _, msg = validar_espesor_mop(capa, espesor, trafico_ee_actual)
    alertas_espesores.append(msg)

# MANTENER: Agregamos la columna de alerta para el editor visual de Streamlit
st.session_state.df_estructura["Validación MOP"] = alertas_espesores

tabla_estructurada = st.data_editor(
    st.session_state.df_estructura,
    disabled=["Material / Capa", "Coef. Estr.", "Coef. Drenaje", "Total (mm)", "Validación MOP"],
    hide_index=True,
    width="stretch",  #"content",
    height=altura_estructura,
    column_config={
        "Material / Capa": st.column_config.TextColumn(width=220),
        "Coef. Estr.": st.column_config.NumberColumn(width=100, format="%.2f"),
        "Espesor (mm)": st.column_config.NumberColumn(width=110, format="%d", min_value=0, step=5),
        "Coef. Drenaje": st.column_config.NumberColumn(width=110, format="%.1f"),
        "Total (mm)": st.column_config.NumberColumn(width=100, format="%.2f"),
        "Validación MOP": st.column_config.TextColumn(width=160), # Nueva columna informativa
    },
    key="editor_estructura"
)

# 🔄 CÁLCULO REACTIVO AUTOMÁTICO (Sin congelamiento)
# Verificamos si hubo cambios en los espesores para recalcular la fila
drenaje_calculo = tabla_estructurada["Coef. Drenaje"].fillna(1.0)
totales_calculados = tabla_estructurada["Coef. Estr."] * tabla_estructurada["Espesor (mm)"] * drenaje_calculo

# Si el cálculo actual difiere de lo guardado, actualizamos y disparamos un solo reinicio controlado
if not tabla_estructurada["Total (mm)"].equals(totales_calculados):
    st.session_state.df_estructura["Espesor (mm)"] = tabla_estructurada["Espesor (mm)"]
    st.session_state.df_estructura["Total (mm)"] = totales_calculados
    st.rerun()  # Ejecuta una sola vez para plasmar el nuevo resultado en pantalla
    
# ==========================================
# 4. TABLA: RESULTADOS (POR FILAS), VERIFICACIÓN DE CAPAS
# ==========================================
col_01, col_02 = st.columns(2)
with col_01:
    st.markdown("### Verificación de Capas y Números Estructurales")

# --- EXTRACCIÓN DE ENTRADAS PARA MCV3 Y AASHTO ---
# Tráfico Solicitante (EE) en millones (ej: 1.43)
ee_millones = t_diseno_calculado/1_000_000   #tabla_entrada.loc[1, "Subrasante"] / 1_000_000 

# Módulo Resiliente de la Base Granular (Fila 8 de la columna Base)
mr_base_mpa = tabla_entrada.loc[8, "Base"]

# Módulo Resiliente de la Base Granular (Fila 8 de la columna Base)
# mr_subrasante_mpa = tabla_entrada.loc[8, "Subrasante"]

# TMAPA dinámico de la estación seleccionada
tmapa_estacion = tmapa_proyecto  

# ==========================================================
# 🟢 RUTINA BLINDADA: REFACTORIZACIÓN MATEMÁTICA MULTICAPA (AASHTO 93)
# ==========================================================
# 1. Purificamos y aseguramos que los MR de la Tabla 1 sean números reales
mr_sub_mpa = limpiar_numero(tabla_entrada.loc[8, "Subrasante"])
mr_base_mpa = limpiar_numero(tabla_entrada.loc[8, "Base"])
mr_subbase_mpa = limpiar_numero(tabla_entrada.loc[8, "Subbase"])

with col_02:
    st.write(f"MResiliente= {mr_sub_mpa: .1f} MPa") 

# --- CALCULO SIMULTÁNEO DE LAS DOS ECUACIONES ---
# A. Vía Manual de Carreteras (Fórmula macro VBA)
nea_requerido_mop_vba = calcular_sn_requerido_mop(t_diseno_calculado, zr_val, so_val, mr_base_mpa, pi, pf)
ne_t_requerido_mop = calcular_sn_requerido_mop(t_diseno_calculado, zr_val, so_val, mr_sub_mpa, pi, pf)

# B. Vía AASHTO 93 Original Americana (Fórmula en Pulgadas/PSI)
nea_requerido_aashto_us = calcular_sn_requerido_aashto(t_diseno_calculado, zr_val, so_val, delta_PSI, mr_base_mpa)
ne_t_requerido_aashto = calcular_sn_requerido_aashto(t_diseno_calculado, zr_val, so_val, delta_PSI, mr_sub_mpa)

# C. Criterio de Control por Clima (Spline Cúbico)
nea_mcv3 = calcular_ne_mcv3(ee_millones, mr_sub_mpa, tmapa_estacion)

# --- CRITERIO ENVOLVENTE FINAL ---
# Elegimos el mayor requerimiento de asfalto definitivo considerando el MCV3
nea_requerido = float(nea_mcv3)   #max(float(nea_requerido_mop_vba), float(nea_mcv3))
ne_t_requerido = ne_t_requerido_mop  # Base de control nacional MOP

# 2. Requerido Total sobre la Subrasante (Usa MR Subrasante numérico)
# Multiplicamos por 10 al final de forma segura si tu modelo lo requiere en mm
#ne_t_requerido = calcular_sn_requerido_aashto(t_diseno_calculado, zr_val, so_val, delta_PSI, mr_sub_mpa) * 10

# 3. Requerido sobre la Subbase (Usa MR Subbase numérico)
#ne_base_asf_requerido = calcular_sn_requerido(t_diseno_calculado, zr_val, so_val, delta_PSI, mr_subbase_mpa) * 10

# 4. Requerido sobre la Base (Usa MR Base numérico) - Teórico AASHTO para asfalto
#nea_requerido_aashto = calcular_sn_requerido(t_diseno_calculado, zr_val, so_val, delta_PSI, mr_base_mpa) * 10

# 5. El requerimiento de asfalto final es el MAYOR entre el teórico AASHTO y el normativo MCV3 chileno
#nea_mcv3 = calcular_ne_mcv3(ee_millones, mr_base_mpa, tmapa_estacion)
#nea_requerido = max(float(nea_requerido_aashto), float(nea_mcv3))

# 6. El requerimiento granular corregido por diferencia de capas estructurales
# Requerimiento granular final por diferencias
ne_gr_requerido = round(ne_t_requerido - nea_requerido, 2)
if ne_gr_requerido < 0: ne_gr_requerido = 0.0

# ==========================================================
# 📊 EXTRACCIÓN Y SUMATORIAS DESDE LA TABLA DE ESTRUCTURACIÓN
# ==========================================================

# 1. Copiamos la columna "Total (mm)" fresca desde el editor de la pantalla
# Nota: `.iloc[0:4]` toma las primeras 4 filas (Micro, Concreto, Binder, Base Asf.)
totales_capas_mm = tabla_estructurada["Total (mm)"]

# 2. Calculamos el NEA Resistente (Suma de las 4 capas asfálticas y conversión a cm)
nea_resistente = round(totales_capas_mm.iloc[0:4].sum(), 2)

# 3. Calculamos el NE Granular Resistente (Suma de las filas 4 y 5: Base y Subbase)
ne_gr_resistente = round(totales_capas_mm.iloc[4:6].sum(), 2)

# 4. Número Estructural Total Resistente Aportado
ne_t_resistente = round(nea_resistente + ne_gr_resistente, 2)

# --- EVALUACIÓN DE CUMPLIMIENTOS ---
status_nea = "Cumple" if nea_resistente >= nea_requerido else "⚠️ No Cumple"
status_ne_gr = "Cumple" if ne_gr_resistente >= ne_gr_requerido else "⚠️ No Cumple"
status_ne_t = "Cumple" if ne_t_resistente >= ne_t_requerido else "⚠️ No Cumple"

# --- MATRIZ DE RESULTADOS COMPARATIVA ---
# --- CONSTRUCCIÓN DE LA TABLA MATRIZ EXTENDIDA ---
filas_resultados = [
    {
        "Parámetro Asfalto": "NE_A Requerido (TMAPA)", 
        "Valor A": nea_requerido, 
        "Parámetro Total / Granular": "NE_T Requerido (MOP)", 
        "Valor B": ne_t_requerido
    },
    {
        "Parámetro Asfalto": f"NEA Resistente ({status_nea})", 
        "Valor A": nea_resistente, 
        "Parámetro Total / Granular": f"NE_T Resistente ({status_ne_t})", 
        "Valor B": ne_t_resistente
    },
    {
        "Parámetro Asfalto": "   ↳ Ref. MCV3 (Base)", 
        "Valor A": round(nea_mcv3, 2), 
        "Parámetro Total / Granular": "NE_T Req. (AASHTO US)", 
        "Valor B": ne_t_requerido_aashto  # 🟢 Comparativa Total
    },
    {
        "Parámetro Asfalto": "   ↳ Ref. MCV3 (Base)", 
        "Valor A": round(nea_requerido_mop_vba, 2), # 🟢 Ecuación Chilena
        "Parámetro Total / Granular": "NEgr Requerido", 
        "Valor B": ne_gr_requerido
    },
    {
        "Parámetro Asfalto": "   ↳ Ref. AASHTO 93 (Base)", 
        "Valor A": round(nea_requerido_aashto_us, 2), # 🟢 Ecuación Americana
        "Parámetro Total / Granular": f"NEgr Resistente ({status_ne_gr})", 
        "Valor B": ne_gr_resistente
    }
      #,
    #{
    #    "Parámetro Asfalto": f"Espesor granular ({status_espesor})", 
    #    "Valor A": None, 
    #    "Parámetro Total / Granular": "", 
    #    "Valor B": None
    #}
]

#Tabla de Resultados
df_resultados = pd.DataFrame(filas_resultados)

# Encabezado (~38px) + ~35px por cada fila + padding extra, si la altura cambia dinámicamente
altura_calculada = 38 + (len(df_resultados) * 35) + 5

# Renderizado estable de la tabla final
tabla_final = st.data_editor(
    df_resultados,
    disabled=True,
    hide_index=True,
    width="content",
    height=altura_calculada, # "auto" <--- Adapta la altura automáticamente a todas las filas o altura_calculada,  # <--- Altura dinámica calculada
    column_config={
        "Parámetro Asfalto": st.column_config.TextColumn(width=220),
        "Valor A": st.column_config.NumberColumn(width=100, format="%.2f", alignment="center"),
        "Parámetro Total / Granular": st.column_config.TextColumn(width=210),
        "Valor B": st.column_config.NumberColumn(width=100, format="%.2f", alignment="center"),
    },
    key="editor_resultados_finales"
)

# ==========================================
# 5. BLOQUE DE EXPORTACIÓN
# ==========================================
# Código para el Bloque de Descargas
st.markdown("---")
st.markdown("### 💾 Exportar Memoria de Cálculo")

# --- CONSTRUCCIÓN DEL PDF CON REPORTLAB ---
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def generar_pdf_vial():
    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=14, leading=20, textColor=colors.HexColor("#1A365D"))
    subtitle_style = ParagraphStyle('SubTitleStyle', parent=styles['Heading3'], fontSize=11, leading=15, spaceBefore=8, textColor=colors.HexColor("#2B6CB0"))
    
    # Cabecera del PDF
    story.append(Paragraph(f"MEMORIA DE CÁLCULO: DISEÑO DE PAVIMENTOS FLEXIBLES", title_style))
    story.append(Paragraph(f"<b>Camino:</b> {nombre_camino} | <b>Sector:</b> {nombre_sector}", styles['Normal']))
    story.append(Paragraph("Método AASHTO 93 / Manual de Carreteras Volumen 3", styles['Normal']))
    story.append(Spacer(1, 10))
    
    # ==========================================
    # 1. TABLA: DATOS DE ENTRADA (EN EL PDF)
    # ==========================================
    story.append(Paragraph("1. Datos de Entrada", subtitle_style))
    
    df_pdf_ent = st.session_state.df_entrada.copy()
    for idx, fila in df_pdf_ent.iterrows():
        val = fila["Subrasante"]
        if pd.notna(val) and val != "":
            # 🟢 SOLUCIÓN: Convertimos el texto formateado a número puro temporalmente
            val_float = float(str(val).replace(",", ""))
            
            if idx == 1 or idx == 9: # Filas de Tráfico 
                df_pdf_ent.at[idx, "Subrasante"] = f"{val_float:,.0f}"
            elif idx == 4: # Error estándar
                df_pdf_ent.at[idx, "Subrasante"] = f"{val_float:,.3f}"
            else:
                df_pdf_ent.at[idx, "Subrasante"] = f"{val_float:,.2f}"

    df_pdf_ent = df_pdf_ent.fillna("")
    data_entrada = [df_pdf_ent.columns.tolist()] + df_pdf_ent.values.tolist()
    
    t1 = Table(data_entrada, colWidths=[210, 100, 100, 100])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1A365D")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('ALIGN', (1,0), (-1,-1), 'CENTER'),
    ]))
    story.append(t1)
    story.append(Spacer(1, 10))
    
    # ==========================================
    # 2. TABLA: ESTRUCTURACIÓN
    # ==========================================
    story.append(Paragraph("2. Estructuración y Espesores de Diseño", subtitle_style))
    
    df_pdf_est = st.session_state.df_estructura.copy()
    if "Total (mm)" in df_pdf_est.columns:
        df_pdf_est["Total (mm)"] = df_pdf_est["Total (mm)"].round(2)
        
    if "Validación MOP" in df_pdf_est.columns:
        df_pdf_est["Validación MOP"] = df_pdf_est["Validación MOP"].replace("No considerada", "")
        
    df_pdf_est = df_pdf_est.fillna("")
    data_est = [df_pdf_est.columns.tolist()] + df_pdf_est.values.tolist()
    
    t2 = Table(data_est, colWidths=[180, 65, 75, 75, 65, 50])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#2B6CB0")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('ALIGN', (1,0), (-1,-1), 'CENTER'),
    ]))
    story.append(t2)
    story.append(Spacer(1, 10))

    # ==========================================
    # 3. 🟢 NUEVA TABLA: VERIFICACIÓN FINAL (AASHTO / MCV3)
    # ==========================================
    story.append(Paragraph("3. Verificación de Capas y Números Estructurales", subtitle_style))
    
    # Clonamos la tabla de resultados calculada en la interfaz de Streamlit
    df_pdf_res = df_resultados.copy()
    df_pdf_res = df_pdf_res.fillna("") # Limpia celdas vacías estéticas
    
    data_res = [df_pdf_res.columns.tolist()] + df_pdf_res.values.tolist()
    
    t3 = Table(data_res, colWidths=[210, 45, 210, 45])
    t3.setStyle(TableStyle([
        # Color Gris Oscuro elegante de Control de Calidad
        ('BACKGROUND', (0,0), (1,0), colors.HexColor("#4A5568")), 
        ('BACKGROUND', (2,0), (3,0), colors.HexColor("#4A5568")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('ALIGN', (1,0), (1,-1), 'CENTER'), # Centra los valores de la columna A
        ('ALIGN', (3,0), (3,-1), 'CENTER'), # Centra los valores de la columna B
    ]))
    story.append(t3)
    
    # Compilación final del documento
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

        
#---------------------------------------------------------------------------
#--------------- GENERA ARCHIVO EXCEL --------------------------------------
#---------------------------------------------------------------------------


def generar_excel_vial(df_entrada, df_materiales, df_estructura, df_resultados, camino="No definido", sector="No definido"):
    from io import BytesIO
    import pandas as pd
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        sheet_name = "Diseño Estructural"
        
        # 1. Limpiar celdas con NaN o vacías para evitar errores de formato en Excel
        df_entrada = df_entrada.fillna("")
        df_materiales = df_materiales.fillna("")
        df_estructura = df_estructura.fillna("")
        df_resultados = df_resultados.fillna("")
        
        # 2. Insertar los DataFrames en filas específicas (dejando espacio para títulos)
        df_entrada.to_excel(writer, sheet_name=sheet_name, index=False, startrow=4)
        df_materiales.to_excel(writer, sheet_name=sheet_name, index=False, startrow=17)
        df_estructura.to_excel(writer, sheet_name=sheet_name, index=False, startrow=25)
        df_resultados.to_excel(writer, sheet_name=sheet_name, index=False, startrow=34) # 🟢 Agregamos la tabla final
        
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        
        # --- 3. CABECERA Y TÍTULOS DEL PROYECTO ---
        worksheet.merge_cells('A1:F1')
        worksheet["A1"] = "MEMORIA DE CÁLCULO: DISEÑO DE PAVIMENTOS FLEXIBLES"
        worksheet["A1"].font = Font(size=13, bold=True, color="FFFFFF")
        worksheet["A1"].fill = PatternFill(start_color="1A365D", end_color="1A365D", fill_type="solid")
        worksheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
        worksheet.row_dimensions[1].height = 25
        
        worksheet.merge_cells('A2:F2')
        worksheet["A2"] = f"Camino: {camino}   |   Sector: {sector}"
        worksheet["A2"].font = Font(size=11, bold=True, color="1A365D")
        worksheet["A2"].alignment = Alignment(horizontal="center", vertical="center")
        
        worksheet.merge_cells('A3:F3')
        worksheet["A3"] = "Método AASHTO 93 / Manual de Carreteras Volumen 3"
        worksheet["A3"].font = Font(size=10, italic=True)
        worksheet["A3"].alignment = Alignment(horizontal="center", vertical="center")
        
        # --- 4. TÍTULOS DE CADA SECCIÓN ---
        titulos = {
            4: "1. Datos de Entrada",
            17: "2. Características de Materiales",
            25: "3. Estructuración y Espesores de Diseño",
            34: "4. Verificación de Capas y Números Estructurales"
        }
        
        for row_idx, texto in titulos.items():
            cell = worksheet.cell(row=row_idx, column=1)
            cell.value = texto
            cell.font = Font(size=11, bold=True, color="2B6CB0")
        
        # --- 5. ESTILOS BASE PARA LAS TABLAS ---
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2B6CB0", end_color="2B6CB0", fill_type="solid")
        center_alignment = Alignment(horizontal="center", vertical="center")
        left_alignment = Alignment(horizontal="left", vertical="center")
        thin_border = Border(
            left=Side(style='thin', color="A0A0A0"), right=Side(style='thin', color="A0A0A0"), 
            top=Side(style='thin', color="A0A0A0"), bottom=Side(style='thin', color="A0A0A0")
        )
        
        # 6. Ensanchar las columnas para que el texto respire
        worksheet.column_dimensions['A'].width = 38
        worksheet.column_dimensions['B'].width = 18
        worksheet.column_dimensions['C'].width = 18
        worksheet.column_dimensions['D'].width = 18
        worksheet.column_dimensions['E'].width = 18
        worksheet.column_dimensions['F'].width = 18
        
        # 7. Función interna para "Pintar" los bordes y colores de los DataFrames
        def format_table(start_row_pandas, df):
            start_row_excel = start_row_pandas + 1
            # Formatear Cabeceras Azules
            for col_num in range(1, len(df.columns) + 1):
                cell = worksheet.cell(row=start_row_excel, column=col_num)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_alignment
                cell.border = thin_border
            
            # Formatear Datos Interiores
            for row_num in range(start_row_excel + 1, start_row_excel + 1 + len(df)):
                for col_num in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.border = thin_border
                    # La columna A y C (nombres de capas/parámetros) van a la izquierda, los números al centro
                    if col_num == 1 or col_num == 3: 
                        cell.alignment = left_alignment
                    else:
                        cell.alignment = center_alignment
        
        # Aplicar el pincelazo a las 4 tablas
        format_table(4, df_entrada)
        format_table(17, df_materiales)
        format_table(25, df_estructura)
        format_table(34, df_resultados)
        
    return output.getvalue()

from io import BytesIO
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    """Auxiliar para aplicar color de fondo a las celdas de una tabla Word"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generar_word_vial(df_entrada, df_materiales, df_estructura, df_resultados, camino="No definido", sector="No definido"):
    doc = Document()

    # Configuración de márgenes del documento (2 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- 1. CABECERA Y TÍTULOS ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("MEMORIA DE CÁLCULO: DISEÑO DE PAVIMENTOS FLEXIBLES")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Camino: {camino}   |   Sector: {sector}")
    run_sub.bold = True
    run_sub.font.size = Pt(11)

    p_method = doc.add_paragraph()
    p_method.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_method = p_method.add_run("Método AASHTO 93 / Manual de Carreteras Volumen 3")
    run_method.italic = True
    run_method.font.size = Pt(10)
    
    doc.add_paragraph()  # Espaciador

    # --- 2. FUNCIÓN PARA INSERTAR TABLAS FORMATEADAS ---
    def agregar_tabla_word(df, titulo_seccion):
        # Título de Sección
        h = doc.add_heading(level=2)
        run_h = h.add_run(titulo_seccion)
        run_h.font.size = Pt(12)
        run_h.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        h.paragraph_format.space_after = Pt(4)

        df_clean = df.fillna("")
        rows, cols = df_clean.shape[0] + 1, df_clean.shape[1]
        
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Fila de Encabezados (Azul #2B6CB0)
        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(df_clean.columns):
            cell = hdr_cells[col_idx]
            cell.text = str(col_name)
            set_cell_background(cell, "2B6CB0")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Filas de Datos
        for row_idx, row_data in enumerate(df_clean.itertuples(index=False)):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, val in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(8.5)
                    # Alineación: Texto a la izquierda (col 0 y 2), Números al centro
                    if col_idx in [0, 2]:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espaciador entre tablas

    # --- 3. CONSTRUCCIÓN DE LAS 4 TABLAS ---
    agregar_tabla_word(df_entrada, "1. Datos de Entrada")
    agregar_tabla_word(df_materiales, "2. Características de Materiales")
    agregar_tabla_word(df_estructura, "3. Estructuración y Espesores de Diseño")
    agregar_tabla_word(df_resultados, "4. Verificación de Capas y Números Estructurales")

    # --- 4. RETORNO DE BUFFER ---
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# --- BOTONES DE INTERFAZ CORREGIDOS (AL FINAL DEL ARCHIVO) ---
col_xlsx, col_pdf, col_docx = st.columns(3)

with col_xlsx:
    # 1. Botón visual para procesar el archivo sin congelar
    if st.button("📊 Preparar Planilla Excel (.xlsx)", width="stretch"):
        excel_data = generar_excel_vial(
            st.session_state.df_entrada, 
            st.session_state.df_materiales, 
            st.session_state.df_estructura, 
            df_resultados,
            camino=nombre_camino,      # 🟢 Inyectamos el nombre del camino
            sector=nombre_sector       # 🟢 Inyectamos el nombre del sector
        )
        # 2. Una vez calculado, se habilita la descarga real al instante
        st.download_button(
            label="⬇️ Hacer clic aquí para Guardar Excel",
            data=excel_data,
            file_name="Memoria_Diseno_Pavimentos_Flexibles.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            width="stretch"
        )

with col_pdf:
    # 1. Botón visual para compilar el PDF con ReportLab
    if st.button("📄 Preparar Memoria PDF (.pdf)", width="stretch"):
        pdf_data = generar_pdf_vial()
        # 2. Una vez compilado, se habilita la descarga real al instante
        st.download_button(
            label="⬇️ Hacer clic aquí para Guardar PDF",
            data=pdf_data,
            file_name="Memoria_Diseno_Pavimentos_Flexibles.pdf",
            mime="application/pdf",
            width="stretch"
        )

with col_docx:
    if st.button("📝 Preparar Word (.docx)", width="stretch"):
        docx_data = generar_word_vial(
            st.session_state.df_entrada, 
            st.session_state.df_materiales, 
            st.session_state.df_estructura, 
            df_resultados,
            camino=nombre_camino,
            sector=nombre_sector
        )
        st.download_button(
            label="⬇️ Descargar Word",
            data=docx_data,
            file_name="Memoria_Diseno_Pavimentos_Flexibles.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width="stretch"
        )

